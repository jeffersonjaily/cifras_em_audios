# exporter.py

from docx import Document
from fpdf import FPDF

def exportar_para_txt(texto, caminho):
    """Salva o texto em um arquivo .txt com codificação UTF-8."""
    with open(caminho, 'w', encoding='utf-8') as f:
        f.write(texto)

def exportar_para_docx(texto, caminho):
    """Salva o texto em um arquivo .docx usando uma fonte monoespaçada."""
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.name = 'Courier New'
    doc.save(caminho)

def exportar_para_pdf(texto, caminho):
    """Salva o texto em um arquivo .pdf usando uma fonte monoespaçada."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Courier", size=10)
    pdf.multi_cell(0, 5, txt=texto)
    pdf.output(caminho)